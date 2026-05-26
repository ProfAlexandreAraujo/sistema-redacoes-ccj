"""
utils.py — Utilitários: leitura/escrita de docx, salvamento de sessão
Sistema de Redações — CCJ CMRJ
"""

import json
import re
import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from harmonizer import Emenda, TipoEmenda, StatusEmenda, ResultadoHarmonizacao


# ─────────────────────────────────────────────────────────────────────────────
# LEITURA DE DOCX / TXT
# ─────────────────────────────────────────────────────────────────────────────

def ler_docx(arquivo_bytes: bytes) -> str:
    """Extrai texto de um arquivo .docx preservando estrutura básica."""
    doc = Document(BytesIO(arquivo_bytes))
    paragrafos = []
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto:
            paragrafos.append(texto)
    return '\n'.join(paragrafos)


def ler_txt(arquivo_bytes: bytes) -> str:
    """Decodifica arquivo de texto."""
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            return arquivo_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return arquivo_bytes.decode('utf-8', errors='replace')


def ler_pdf(arquivo_bytes: bytes) -> str:
    """Extrai texto de arquivo PDF usando pdfplumber.
    Remove cabeçalhos e rodapés típicos de impressões web da CMRJ.
    """
    import re as _re
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber não instalado. Execute: pip install pdfplumber")

    paginas = []
    with pdfplumber.open(BytesIO(arquivo_bytes)) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ''
            # Remove cabeçalhos/rodapés de PDFs impressos do site da CMRJ
            texto = _re.sub(
                r'\d{2}/\d{2}/\d{4},?\s*\d{2}:\d{2}\s+Projeto de Lei.*?\n', '', texto
            )
            texto = _re.sub(
                r'https?://www\.camara\.rio/[^\n]+\n?', '', texto
            )
            paginas.append(texto)

    texto_completo = '\n'.join(paginas)

    # Truncar no ponto mais cedo entre todos os marcadores de parada.
    # Usa o mínimo (não o primeiro da lista) para garantir que seções como
    # JUSTIFICATIVA (anterior a TRAMITAÇÃO) sejam cortadas corretamente.
    _stop_marcadores = (
        'JUSTIFICATIVA',
        'Texto Original:',
        'LEGISLAÇÃO CITADA',
        'MENSAGEM Nº',
        'TRAMITAÇÃO DO PROJETO',
        'Distribuição =>',
        'Informações Básicas',
    )
    stop_pos = len(texto_completo)
    for marcador in _stop_marcadores:
        idx = texto_completo.find(marcador)
        if idx > 500:           # ignora se aparecer logo no início
            stop_pos = min(stop_pos, idx)
    texto_completo = texto_completo[:stop_pos]

    return texto_completo.strip()


# ─────────────────────────────────────────────────────────────────────────────
# ANÁLISE ESTRUTURAL RÁPIDA
# ─────────────────────────────────────────────────────────────────────────────

def analisar_estrutura(texto: str) -> dict:
    """Retorna contagem de elementos estruturais do projeto.

    Regras baseadas na LC 95/1998, Decreto 12.002/2024 e LC Municipal 48/2000:
    - Artigos: apenas cabeçalhos (início de linha), nunca referências internas
    - Parágrafos: cabeçalhos § Nº ou "Parágrafo único" no início da linha
    - Incisos: algarismos romanos seguidos de travessão (– ou —) ou hífen (-)
    - Alíneas: letras minúsculas seguidas de ) no início da linha (com indentação)
    - Anexos: palavra ANEXO seguida de numeral romano ou arábico no início da linha
    """
    # Artigos: início de linha, seguido de número ordinal (com ou sem º/o)
    # Evita contar referências como "conforme o Art. 3º"
    # O ponto após "Art" é opcional para cobrir erros tipográficos ("Art 14.")
    artigos    = re.findall(
        r'^Art\.?\s*\d+[ºo°]?',
        texto, re.MULTILINE | re.IGNORECASE
    )

    # Parágrafos: início de linha (com possível indentação)
    # Evita contar referências como "previsto no § 1º"
    paragrafos = re.findall(
        r'^\s*(?:§\s*\d+[ºo°]?|Parágrafo\s+único)',
        texto, re.MULTILINE | re.IGNORECASE
    )

    # Incisos: numerais romanos no início da linha (com ou sem indentação)
    # seguidos de travessão (–, —) ou hífen (-), conforme LC 95/98 art. 13
    # Cobre até XX para projetos extensos
    incisos    = re.findall(
        r'^\s*(?:X{0,2}(?:IX|IV|V?I{0,3}))\s*[-–—]',
        texto, re.MULTILINE
    )
    # Filtrar falsos positivos: remover matches vazios (ex: "—" sozinho)
    incisos = [m for m in incisos if re.search(r'[IVX]', m)]

    # Alíneas: letras minúsculas a-z seguidas de ) — LC 95/98 art. 13 III
    alineas    = re.findall(
        r'^\s+[a-z]\)\s',
        texto, re.MULTILINE
    )

    # Anexos: início de linha
    anexos     = re.findall(
        r'^\s*ANEXO\s+[IVX\d]+',
        texto, re.MULTILINE | re.IGNORECASE
    )

    return {
        'artigos':    len(artigos),
        'paragrafos': len(paragrafos),
        'incisos':    len(incisos),
        'alineas':    len(alineas),
        'anexos':     len(anexos),
    }


# ─────────────────────────────────────────────────────────────────────────────
# EXTRAÇÃO DE METADADOS DO PROJETO
# ─────────────────────────────────────────────────────────────────────────────

def extrair_ementa_autor(texto: str) -> tuple[str, str]:
    """Extrai a ementa e a linha de autor do texto original do projeto.

    Retorna: (ementa, linha_autor)
    Procura os padrões comuns da CMRJ:
        EMENTA:
        TEXTO DA EMENTA...
        Autor(es): VEREADOR(A) XXXXX
    """
    ementa = ""
    autor  = ""

    # Ementa: texto entre "EMENTA:" e a próxima seção relevante
    m_em = re.search(
        r'EMENTA:\s*\n([\s\S]*?)(?=\n\s*(?:Autor\(es\)|A\s+C[ÂA]MARA|$))',
        texto, re.IGNORECASE
    )
    if m_em:
        # Colapsa quebras de linha e espaços extras em espaço único
        ementa = ' '.join(m_em.group(1).split())

    # Autor(es): primeira ocorrência
    m_aut = re.search(r'(Autor\(es\)\s*:.*)', texto, re.IGNORECASE)
    if m_aut:
        autor = m_aut.group(1).strip()

    return ementa, autor


# ─────────────────────────────────────────────────────────────────────────────
# SALVAMENTO / CARREGAMENTO DE SESSÃO
# ─────────────────────────────────────────────────────────────────────────────

SAVES_DIR = Path(__file__).parent / "sessoes_salvas"


def salvar_sessao(nome_projeto: str, texto_original: str, emendas: list[Emenda]) -> Path:
    """Salva estado da sessão em JSON."""
    SAVES_DIR.mkdir(exist_ok=True)
    slug = re.sub(r'[^\w]', '_', nome_projeto or 'projeto')[:40]
    ts   = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    caminho = SAVES_DIR / f"{slug}_{ts}.json"
    dados = {
        "nome_projeto":   nome_projeto,
        "texto_original": texto_original,
        "emendas":        [e.to_dict() for e in emendas],
        "salvo_em":       ts,
    }
    with open(caminho, 'w', encoding='utf-8') as f:
        json.dump(dados, f, ensure_ascii=False, indent=2)
    return caminho


def listar_sessoes() -> list[Path]:
    """Lista arquivos de sessão salvos, mais recente primeiro."""
    if not SAVES_DIR.exists():
        return []
    return sorted(SAVES_DIR.glob("*.json"), reverse=True)


def carregar_sessao(caminho: Path) -> tuple[str, str, list[Emenda]]:
    """Carrega sessão salva. Retorna (nome_projeto, texto_original, emendas)."""
    with open(caminho, 'r', encoding='utf-8') as f:
        dados = json.load(f)
    emendas = [Emenda.from_dict(d) for d in dados.get("emendas", [])]
    return dados.get("nome_projeto", ""), dados.get("texto_original", ""), emendas


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS INTERNOS PARA EXPORTAÇÃO DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _aplicar_sufixo_a(nome: str) -> str:
    """
    Insere o sufixo -A obrigatório no número do projeto.
    Exemplos:
        'PLC 92/2025'                      → 'PLC 92-A/2025'
        'PLC 92/2025 — AEIU Praça XI'     → 'PLC 92-A/2025 — AEIU Praça XI'
        'PLC 92 2025'  (sem barra)         → 'PLC 92 2025'  (não altera)
    """
    return re.sub(r'(\d+)\s*(/\s*\d{4})', r'\1-A\2', nome)


def _remover_bordas_tabela(tabela) -> None:
    """Remove todas as bordas visíveis de uma tabela python-docx."""
    # Imports lazy: evita falha de importação no nível de módulo em ambientes
    # onde as APIs internas do python-docx estão em caminhos diferentes.
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    for row in tabela.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = _OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                tag = _OxmlElement(f'w:{edge}')
                tag.set(_qn('w:val'),   'none')
                tag.set(_qn('w:sz'),    '0')
                tag.set(_qn('w:space'), '0')
                tag.set(_qn('w:color'), 'auto')
                tcBorders.append(tag)
            tcPr.append(tcBorders)


# Meses em português para o fecho
_MESES_PT = {
    1: 'janeiro',  2: 'fevereiro', 3: 'março',    4: 'abril',
    5: 'maio',     6: 'junho',     7: 'julho',     8: 'agosto',
    9: 'setembro', 10: 'outubro',  11: 'novembro', 12: 'dezembro',
}

# Padrões para formatação especial do corpo
_CAMARA_RE  = re.compile(r'^\s*A\s+C[ÂA]MARA\s+MUNICIPAL', re.IGNORECASE)
_DECRETA_RE = re.compile(r'^\s*D[\s]*E[\s]*C[\s]*R[\s]*E[\s]*T[\s]*A', re.IGNORECASE)
_MARKER_RE  = re.compile(r'\s*\[\[⚠️ CCJ:[^\]]*\]\]', re.UNICODE)


# ─────────────────────────────────────────────────────────────────────────────
# EXPORTAÇÃO PARA DOCX — FORMATO OFICIAL CMRJ/CCJ
# ─────────────────────────────────────────────────────────────────────────────

def exportar_redacao_final_docx(
    texto: str,
    nome_projeto: str,
    avisos: list[str],
    erros: list[str],
    alertas_absurdos: list[str] = None,
    mapa: dict = None,
    log: list[str] = None,
    tipo_redacao: str = "Redação Final",
    prosseguir_com_alerta_sec_2: bool = False,
    ementa: str = "",
    autor: str = "",
) -> bytes:
    """
    Gera arquivo .docx formatado segundo o modelo oficial CMRJ/CCJ.

    Estrutura do documento:
      1. Título (REDAÇÃO FINAL / REDAÇÃO DO VENCIDO / RASCUNHO)
      2. Número do projeto
      3. Ementa (se disponível)
      4. Autor(es) (se disponível)
      5. Corpo do texto harmonizado
      6. Fecho: Sala da Comissão + assinaturas dos 3 vereadores
      7. Anexo de avisos/erros (se houver) — página separada
    """
    # ── Normalizar listas ────────────────────────────────────────────────────
    _alertas_norm = alertas_absurdos or []
    _erros_norm   = erros or []
    mapa          = mapa  or {}
    log           = log   or []

    tem_sec_2   = bool(_erros_norm or _alertas_norm)
    eh_rascunho = tem_sec_2 and not prosseguir_com_alerta_sec_2

    # ── Criar documento ──────────────────────────────────────────────────────
    doc = Document()

    # Configuração da página: A4, margens 2,5 cm
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # Estilo padrão: Times New Roman 10 pt, sem espaçamento extra
    normal = doc.styles['Normal']
    normal.font.name  = 'Times New Roman'
    normal.font.size  = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # ── Helpers locais ───────────────────────────────────────────────────────
    def _blank():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    def _para(text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              bold: bool = False, size_pt: float = 10,
              underline: bool = False, italic: bool = False,
              color: RGBColor = None,
              sp_before: float = 0, sp_after: float = 0) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sp_before)
        p.paragraph_format.space_after  = Pt(sp_after)
        if text:
            r = p.add_run(text)
            r.bold      = bold
            r.italic    = italic
            r.underline = underline
            r.font.name = 'Times New Roman'
            r.font.size = Pt(size_pt)
            if color:
                r.font.color.rgb = color

    def _footer_para(text: str, bold: bool = False):
        _para(text, align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=bold, size_pt=10, sp_before=6, sp_after=6)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. TÍTULO
    # ══════════════════════════════════════════════════════════════════════════
    _COR_ALERTA = RGBColor(0xC0, 0x00, 0x00)

    if eh_rascunho:
        titulo_doc = "RASCUNHO DE TRABALHO — NÃO É REDAÇÃO FINAL"
        _para(titulo_doc,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, size_pt=10,
              color=_COR_ALERTA)
        _para(
            "⚠ RASCUNHO — existem alertas de §2º (absurdo manifesto ou erro crítico) "
            "que exigem avaliação antes da publicação. "
            "Confirme ciência na aba 5 para exportar como Redação Final.",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, italic=True, size_pt=9,
            color=_COR_ALERTA,
        )
    else:
        titulo_doc = tipo_redacao.upper()
        _para(titulo_doc,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, underline=True, size_pt=10)
        if tem_sec_2:
            _para(
                "⚠ ALERTA CRÍTICO PENDENTE — ART. 250, §2º RI — "
                "Relator tomou ciência e optou por prosseguir. Ver Anexo de Avisos.",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True, size_pt=9, color=_COR_ALERTA,
            )

    # Três linhas em branco após título (padrão modelo)
    _blank(); _blank(); _blank()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. NÚMERO DO PROJETO
    # ══════════════════════════════════════════════════════════════════════════
    if nome_projeto:
        nome_doc = nome_projeto if eh_rascunho else _aplicar_sufixo_a(nome_projeto)
        _para(nome_doc.upper(),
              align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, size_pt=10)

    _blank()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. EMENTA
    # ══════════════════════════════════════════════════════════════════════════
    _para("EMENTA:",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          bold=True, size_pt=10)

    if ementa:
        # Ementa em tabela sem bordas (padrão CMRJ)
        tbl_em = doc.add_table(rows=1, cols=1)
        _remover_bordas_tabela(tbl_em)
        cell_em = tbl_em.cell(0, 0)
        p_em    = cell_em.paragraphs[0]
        p_em.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_em = p_em.add_run(ementa)
        r_em.font.name = 'Times New Roman'
        r_em.font.size = Pt(10)
    else:
        # Espaço reservado para preenchimento manual
        _blank(); _blank(); _blank()

    _blank(); _blank(); _blank()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. AUTOR(ES)
    # ══════════════════════════════════════════════════════════════════════════
    if autor:
        _para(autor,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              bold=True, size_pt=10)

    _blank(); _blank()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. CORPO DO TEXTO HARMONIZADO
    # ══════════════════════════════════════════════════════════════════════════
    texto_limpo = _MARKER_RE.sub('', texto)

    for linha in texto_limpo.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        if _CAMARA_RE.match(linha):
            # "A CÂMARA MUNICIPAL DO RIO DE JANEIRO" — justificado, negrito
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(linha)
            r.bold      = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

        elif _DECRETA_RE.match(linha):
            # "DECRETA:" / "D E C R E T A :" — alinhado à direita, negrito
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(linha)
            r.bold      = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

        else:
            # Demais linhas — justificado, sem negrito
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(linha)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════════
    # 6. FECHO — Sala da Comissão + Assinaturas
    # ══════════════════════════════════════════════════════════════════════════
    hoje    = datetime.date.today()
    data_pt = f"{hoje.day} de {_MESES_PT[hoje.month]} de {hoje.year}"

    _blank()
    _footer_para(f"Sala da Comissão, {data_pt}.")
    _footer_para("")
    _footer_para("")
    _footer_para("Vereador Átila Nunes")
    _footer_para("Presidente")

    # Tabela 2×2 sem bordas para Dr. Gilberto | Inaldo Silva
    tbl_sig = doc.add_table(rows=2, cols=2)
    _remover_bordas_tabela(tbl_sig)

    _sig_dados = [
        (0, 0, "Vereador Dr. Gilberto"),
        (0, 1, "Vereador Inaldo Silva"),
        (1, 0, "Vice-presidente"),
        (1, 1, "Vogal"),
    ]
    for row_i, col_i, text in _sig_dados:
        cell = tbl_sig.cell(row_i, col_i)
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════════════════
    # 7. ANEXO DE AVISOS E ALERTAS (página separada, se houver)
    # ══════════════════════════════════════════════════════════════════════════
    if avisos or erros or _alertas_norm:
        doc.add_page_break()
        doc.add_heading('ANEXO — AVISOS E ALERTAS DA CCJ', level=2)
        doc.add_paragraph(
            "Os seguintes pontos foram identificados durante a harmonização e devem ser "
            "avaliados pelo relator antes da publicação da redação final."
        )
        doc.add_paragraph()

        if _alertas_norm:
            h3 = doc.add_heading(
                '🔴 ABSURDO MANIFESTO — REABERTURA DA DISCUSSÃO (art. 250, §2º, RI)',
                level=3,
            )
            h3.runs[0].font.color.rgb = _COR_ALERTA
            doc.add_paragraph(
                "Os itens abaixo configuram incoerência notória, contradição evidente ou "
                "manifesto absurdo capaz de gerar dúvida quanto à vontade legislativa. "
                "Nos termos do art. 250, §2º, do Regimento Interno, a CCJ deverá "
                "eximir-se de oferecer Redação Final e propor, em parecer, a reabertura "
                "da discussão quanto aos aspectos indicados."
            )
            for al in _alertas_norm:
                doc.add_paragraph(f"🔴  {al}", style='List Bullet')

        if erros:
            h3 = doc.add_heading(
                '🚨 ERROS CRÍTICOS — REABERTURA DA DISCUSSÃO', level=3
            )
            h3.runs[0].font.color.rgb = _COR_ALERTA
            doc.add_paragraph(
                "Os itens abaixo envolvem contradição entre emendas aprovadas. "
                "A CCJ deve propor reabertura da discussão (art. 250, §2º, RI)."
            )
            for e in erros:
                doc.add_paragraph(f"🚨  {e}", style='List Bullet')

        if avisos:
            doc.add_heading('⚠️ Avisos redacionais (art. 250, §1º, RI)', level=3)
            doc.add_paragraph(
                "Os itens abaixo compreendem correções de linguagem incorporadas à minuta "
                "(registradas no log para formalização pela CCJ) ou impropriedades apontadas "
                "sem alteração do texto aprovado — em ambos os casos, observada a "
                "formalização prevista no art. 250, §1º, do Regimento Interno."
            )
            for a in avisos:
                doc.add_paragraph(f"⚠  {a}", style='List Bullet')

    if mapa:
        doc.add_paragraph()
        doc.add_heading('Mapa de Renumeração', level=3)
        for orig, novo in mapa.items():
            doc.add_paragraph(f"{orig}  →  {novo}", style='List Bullet')

    # Log de alterações (com registro de override §2º se aplicável)
    log_final = list(log)
    if tem_sec_2 and prosseguir_com_alerta_sec_2:
        log_final.append(
            f"OVERRIDE-HUMANO / Art. 250, §2º RI — Relator tomou ciência dos alertas "
            f"({len(_erros_norm)} erro(s) crítico(s), {len(_alertas_norm)} absurdo(s) manifesto(s)) "
            f"e optou por prosseguir com a Redação Final em {hoje.strftime('%d/%m/%Y')}."
        )
    if log_final:
        doc.add_paragraph()
        doc.add_heading('Log de Alterações Aplicadas', level=3)
        for item in log_final:
            doc.add_paragraph(f"• {item}")

    # ── Serializar ──────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def exportar_relatorio_problemas_txt(
    nome_projeto: str,
    avisos: list[str],
    erros: list[str],
    mapa: dict,
    alertas_absurdos: list[str] = None,
) -> str:
    """Gera relatório de problemas em texto simples."""
    linhas = [
        "=" * 60,
        "RELATÓRIO DE PROBLEMAS — SISTEMA DE REDAÇÕES CCJ CMRJ",
        "=" * 60,
        f"Projeto: {nome_projeto or 'Não informado'}",
        f"Data:    {datetime.date.today().strftime('%d/%m/%Y')}",
        "",
    ]
    alertas_absurdos = alertas_absurdos or []
    if alertas_absurdos:
        linhas += ["", "🔴 ABSURDO MANIFESTO — REABERTURA DA DISCUSSÃO (art. 250, §2º, RI):", "-" * 50]
        linhas += [f"  🔴 {al}" for al in alertas_absurdos]
    if erros:
        linhas += ["", "🚨 ERROS CRÍTICOS (podem exigir reabertura — art. 250 §2º RI):", "-" * 50]
        linhas += [f"  🚨 {e}" for e in erros]
    if avisos:
        linhas += ["", "⚠ AVISOS REDACIONAIS (art. 250 §1º RI — corrigíveis mediante ofício):", "-" * 50]
        linhas += [f"  ⚠  {a}" for a in avisos]
    if mapa:
        linhas += ["", "MAPA DE RENUMERAÇÃO:", "-" * 50]
        linhas += [f"  {orig}  →  {novo}" for orig, novo in mapa.items()]
    return '\n'.join(linhas)
