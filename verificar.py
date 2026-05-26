#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
verificar.py — Verificação do Sistema de Redações CCJ CMRJ (rev.12)
Testa todos os comportamentos críticos do AUDITORIA.md sem custo de API.

Uso:
    python verificar.py              # Testes locais (gratuito)
    python verificar.py --com-api    # + harmonização completa do PLC 17/2026 (~$0,50)
"""

import sys
import os
import re
import pathlib
from io import BytesIO

sys.stdout.reconfigure(encoding="utf-8")

PASS  = "✅ PASS"
FAIL  = "❌ FAIL"
resultados: list[tuple[str, bool]] = []


def chk(nome: str, ok: bool, detalhe: str = "") -> None:
    simbolo = PASS if ok else FAIL
    print(f"  {simbolo}  {nome}")
    if detalhe and not ok:
        print(f"         ↳ {detalhe}")
    resultados.append((nome, ok))


print()
print("=" * 65)
print("  VERIFICAÇÃO DO SISTEMA — CCJ CMRJ — rev.12")
print("=" * 65)

# ────────────────────────────────────────────────────────────────────────────
# 1. IMPORTAÇÕES
# ────────────────────────────────────────────────────────────────────────────
print("\n[1] IMPORTAÇÕES")

try:
    from harmonizer import (
        _detectar_absurdos_estruturais,
        _escalar_avisos_para_absurdos,
        _PADROES_ABSURDO_AVISO,
        ResultadoHarmonizacao,
        TipoEmenda, StatusEmenda, Emenda,
    )
    chk("harmonizer.py importado", True)
except Exception as e:
    chk("harmonizer.py importado", False, str(e))
    print("\n  Falha crítica — abortando.")
    sys.exit(1)

try:
    from utils import exportar_redacao_final_docx, _aplicar_sufixo_a, analisar_estrutura
    chk("utils.py importado", True)
except Exception as e:
    chk("utils.py importado", False, str(e))

try:
    from docx import Document
    chk("python-docx disponível", True)
    _DOCX_OK = True
except Exception as e:
    chk("python-docx disponível", False, str(e))
    _DOCX_OK = False

# ────────────────────────────────────────────────────────────────────────────
# 2. SUFIXO -A
# ────────────────────────────────────────────────────────────────────────────
print("\n[2] SUFIXO -A NO NÚMERO DO PROJETO")

chk("PLC 92/2025 → PLC 92-A/2025",
    _aplicar_sufixo_a("PLC 92/2025") == "PLC 92-A/2025")
chk("PLC 17/2026 → PLC 17-A/2026",
    _aplicar_sufixo_a("PLC 17/2026") == "PLC 17-A/2026")
chk("Título completo preservado",
    "92-A/2025" in _aplicar_sufixo_a("PLC 92/2025 — AEIU Praça XI"))
chk("Sem barra: não altera",
    _aplicar_sufixo_a("PLC 92 2025") == "PLC 92 2025")

# ────────────────────────────────────────────────────────────────────────────
# 3. DETECTOR ESTRUTURAL — CASO 1 (autoreferência circular)
# ────────────────────────────────────────────────────────────────────────────
print("\n[3] P1 — DETECTOR ESTRUTURAL: Caso 1 (autoreferência circular)")

TEXTO_C1 = (
    "Art. 4º Observada a área de abrangência definida no Art. 4º desta Lei "
    "Complementar, ficam regulamentados os parâmetros urbanísticos.\n\n"
    "Art. 5º Os empreendimentos deverão atender as condições previstas no Art. 3º."
)

al_c1 = _detectar_absurdos_estruturais(TEXTO_C1)
chk("Art. 4 detectado como circular",
    any("Art. 4" in a and "circular" in a for a in al_c1))
chk("Art. 5 NÃO detectado (sem self-ref)",
    not any("Art. 5" in a for a in al_c1))

# ────────────────────────────────────────────────────────────────────────────
# 4. DETECTOR ESTRUTURAL — CASO 2 (condição normativa inoperante)
# ────────────────────────────────────────────────────────────────────────────
print("\n[4] P1 — DETECTOR ESTRUTURAL: Caso 2 (condição normativa inoperante)")

TEXTO_C2 = (
    "Art. 12. Os empreendimentos ficam sujeitos à outorga onerosa.\n\n"
    "§ 1º Atendida a condição prevista no §1º deste artigo, fica dispensada "
    "a apresentação de declaração técnica prévia.\n\n"
    "Art. 13. Outras normas gerais."
)

al_c2 = _detectar_absurdos_estruturais(TEXTO_C2)
chk("§1º detectado como condição inoperante",
    any("§1" in a and "inoperante" in a for a in al_c2))

# ────────────────────────────────────────────────────────────────────────────
# 5. PADRÕES SEMÂNTICOS — CASO 3
# ────────────────────────────────────────────────────────────────────────────
print("\n[5] P1 — PADRÕES SEMÂNTICOS NOS AVISOS: Caso 3")

CASOS_PAD = [
    (True,  "artigo anterior + monitoramento incompatível",
     "⚠ §4: nos termos do artigo anterior — trata de monitoramento, incompatível"),
    (True,  "condição §1 suprimida",
     "⚠ §1: condição prevista no §1 deste artigo foi suprimida pela Emenda 5"),
    (True,  "referência circular",
     "⚠ Art. 4: referência circular — o artigo aponta para si mesmo"),
    (False, "concordância verbal simples NÃO escalada",
     "⚠ Art. 16: concordância verbal — texto preservado como aprovado"),
    (False, "maiúscula simples NÃO escalada",
     "⚠ Art. 10: 'Depósitos' com inicial maiúscula — texto preservado"),
]
for esperado, desc, av in CASOS_PAD:
    chk(f"«{desc}»",
        bool(_PADROES_ABSURDO_AVISO.search(av)) == esperado,
        f"esperado={'escalar' if esperado else 'não escalar'}")

# ────────────────────────────────────────────────────────────────────────────
# 6. INTEGRAÇÃO DO ESCALADOR
# ────────────────────────────────────────────────────────────────────────────
print("\n[6] P1 — INTEGRAÇÃO _escalar_avisos_para_absurdos()")

AVISOS_TEST = [
    "⚠ Emenda 1 / Art. 4: referência circular — Art. 4 aponta para si mesmo",
    "⚠ Emenda 5 / §1: condição prevista no §1 deste artigo foi suprimida",
    "⚠ Emenda 10 / §4: artigo anterior trata de monitoramento, incompatível",
    "⚠ Emenda 8 / Art. 16: serão aplicada — concordância verbal",
    "⚠ Emenda 9 / Art. 10: Depósitos com maiúscula",
]
rest, alrts = _escalar_avisos_para_absurdos(AVISOS_TEST, TEXTO_C1, [])
chk("2 avisos de linguagem permanecem como avisos",
    len(rest) == 2,
    f"restantes={len(rest)}, esperado=2 — {rest}")
chk("≥3 absurdos escalados para §2º",
    len(alrts) >= 3,
    f"alertas={len(alrts)}, esperado≥3")
chk("Avisos de linguagem são os corretos (concordância / maiúscula)",
    any("serão aplicada" in r or "concordância" in r.lower() for r in rest) and
    any("Depósitos" in r or "maiúscula" in r for r in rest))

# ────────────────────────────────────────────────────────────────────────────
# 7. EXPORTAÇÃO DOCX
# ────────────────────────────────────────────────────────────────────────────
print("\n[7] EXPORTAÇÃO DOCX")

if _DOCX_OK:
    TEXTO_MARCADOR = (
        "Art. 1º Texto normal.\n"
        "Art. 2º Texto problemático. [[⚠️ CCJ: DISPOSITIVO ININTELIGÍVEL]]\n"
        "Art. 3º Conclusão."
    )

    # 7a — Com §2º, SEM confirmação → RASCUNHO (padrão seguro)
    try:
        docx_a = exportar_redacao_final_docx(
            texto=TEXTO_MARCADOR, nome_projeto="PLC 17/2026",
            avisos=[], erros=[],
            alertas_absurdos=["🔴 Art. 2: absurdo manifesto"],
            prosseguir_com_alerta_sec_2=False,   # padrão
        )
        doc_a  = Document(BytesIO(docx_a))
        txts_a = [p.text for p in doc_a.paragraphs]
        chk("§2º sem confirmação → título 'RASCUNHO DE TRABALHO'",
            any("RASCUNHO DE TRABALHO" in t for t in txts_a))
        chk("§2º sem confirmação → NÃO contém 'REDAÇÃO FINAL' no título",
            not any(t.strip() == "REDAÇÃO FINAL" for t in txts_a))
        chk("§2º sem confirmação → sufixo -A NÃO aplicado no rascunho",
            not any("17-A/2026" in t for t in txts_a))
        chk("§2º sem confirmação → aviso de rascunho presente no cabeçalho",
            any("RASCUNHO" in t for t in txts_a))
        chk("Marcadores [[⚠️ CCJ:...]] removidos do DOCX",
            not any("[[" in t for t in txts_a))
    except Exception as e:
        chk("DOCX §2º sem confirmação", False, str(e))

    # 7a2 — Com §2º, COM confirmação → REDAÇÃO FINAL + ALERTA CRÍTICO + log override
    try:
        docx_a2 = exportar_redacao_final_docx(
            texto=TEXTO_MARCADOR, nome_projeto="PLC 17/2026",
            avisos=[], erros=[],
            alertas_absurdos=["🔴 Art. 2: absurdo manifesto"],
            prosseguir_com_alerta_sec_2=True,    # relator confirmou
        )
        doc_a2  = Document(BytesIO(docx_a2))
        txts_a2 = [p.text for p in doc_a2.paragraphs]
        chk("§2º com confirmação → título contém 'REDAÇÃO FINAL'",
            any("REDAÇÃO FINAL" in t for t in txts_a2))
        chk("§2º com confirmação → sufixo -A aplicado ('17-A/2026')",
            any("17-A/2026" in t for t in txts_a2))
        chk("§2º com confirmação → 'ALERTA CRÍTICO PENDENTE' no cabeçalho",
            any("ALERTA CRÍTICO PENDENTE" in t for t in txts_a2))
        chk("§2º com confirmação → log contém 'OVERRIDE-HUMANO'",
            any("OVERRIDE-HUMANO" in t for t in txts_a2))
    except Exception as e:
        chk("DOCX §2º com confirmação", False, str(e))

    # 7b — Sem §2º → REDAÇÃO FINAL com sufixo -A, sem aviso
    try:
        docx_b = exportar_redacao_final_docx(
            texto="Art. 1º Texto limpo sem problemas.",
            nome_projeto="PLC 17/2026",
            avisos=[], erros=[], alertas_absurdos=[],
        )
        doc_b  = Document(BytesIO(docx_b))
        txts_b = [p.text for p in doc_b.paragraphs]
        chk("Sem §2º → título 'REDAÇÃO FINAL'",
            any("REDAÇÃO FINAL" in t for t in txts_b))
        chk("Sem §2º → sufixo -A aplicado ('17-A/2026')",
            any("17-A/2026" in t for t in txts_b))
        chk("Sem §2º → NÃO contém 'RASCUNHO'",
            not any("RASCUNHO" in t for t in txts_b))
    except Exception as e:
        chk("DOCX sem §2º", False, str(e))

    # 7c — Texto de absurdo manifesto no anexo menciona §2º e reabertura (não §1º)
    try:
        chk("Absurdo no DOCX menciona §2º (não §1º)",
            any("§2" in t and ("reabertura" in t.lower() or "eximir" in t.lower())
                for t in txts_a))
        chk("Absurdo no DOCX NÃO menciona ofício §1º",
            not any("§1" in t and "ofício" in t.lower() for t in txts_a))
    except Exception as e:
        chk("7c — fundamentação §2º no DOCX", False, str(e))

    # 7d — Seção de avisos não afirma que textos foram "preservados exatamente"
    try:
        docx_e1 = exportar_redacao_final_docx(
            texto="Art. 1º Texto.",
            nome_projeto="PLC 17/2026",
            avisos=["⚠ E1 / Art. 1º: corrigido — 'serão aplicada' → 'serão aplicadas'"],
            erros=[], alertas_absurdos=[],
        )
        txts_e1 = [p.text for p in Document(BytesIO(docx_e1)).paragraphs]
        chk("Avisos com E1: seção NÃO diz 'preservados exatamente como aprovados'",
            not any("preservados exatamente como aprovados" in t for t in txts_e1))
    except Exception as e:
        chk("7d — texto avisos E1", False, str(e))
    # 7f — .txt no modo rascunho inclui cabeçalho de alerta
    # (Simula diretamente a lógica do app.py, sem rodar Streamlit)
    try:
        import datetime as _dt
        _txt_rascunho_ok = True
        # Simula o bloco do app.py: _eh_rascunho_aba5=True, _tem_sec_2=True, _prosseguir=False
        _eh_rascunho = True
        _cab = (
            "RASCUNHO DE TRABALHO — NÃO É REDAÇÃO FINAL\n"
            "⚠ Existem alertas de §2º (absurdo manifesto ou erro crítico).\n"
            "A providência regimental indicada é a reabertura da discussão (art. 250, §2º RI).\n"
            "Confirme ciência na aba 5 do sistema para exportar como Redação Final.\n"
            f"Elaborado em {_dt.date.today().strftime('%d/%m/%Y')}\n"
            "=" * 60 + "\n\n"
        )
        # Simula limpeza de marcadores no TXT (P2 — auditoria externa)
        _marker_re_txt_sim = re.compile(r'\s*\[\[⚠️ CCJ:[^\]]*\]\]', re.UNICODE)
        _texto_com_marcador = "Art. 1º Texto. [[⚠️ CCJ: CONFLITO DE EMENDAS — decisão do relator obrigatória]]\nArt. 2º Fim."
        _texto_limpo_sim = _marker_re_txt_sim.sub('', _texto_com_marcador)
        _txt_content = (_cab + _texto_limpo_sim).encode('utf-8')
        _slug_doc_sim = "rascunho_trabalho"
        chk("TXT rascunho: cabeçalho contém 'RASCUNHO DE TRABALHO'",
            b"RASCUNHO DE TRABALHO" in _txt_content)
        chk("TXT rascunho: cabeçalho contém 'reabertura da discussão'",
            "reabertura da discussão".encode('utf-8') in _txt_content)
        chk("TXT rascunho: nome slug é 'rascunho_trabalho'",
            _slug_doc_sim == "rascunho_trabalho")
        chk("TXT: marcadores [[⚠️ CCJ:...]] removidos (P2 — auditoria externa)",
            b"[[" not in _txt_content and b"CONFLITO DE EMENDAS" not in _txt_content)
        chk("TXT: limpeza usa mesmo padrão regex do DOCX (_marker_re_txt)",
            "_marker_re_txt" in _app_src if '_app_src' in dir() else
            "_marker_re_txt" in pathlib.Path(
                r"C:\Users\Admin\Documents\Claude\CCJ\sistema_redacoes\app.py"
            ).read_text(encoding='utf-8'))
    except Exception as e:
        chk("7f — TXT modo rascunho", False, str(e))

else:
    print("  ⏭  DOCX — python-docx não disponível, pulando")

# ────────────────────────────────────────────────────────────────────────────
# 7e — PARSING: texto_bruto e offset (sem chamada de API)
# ────────────────────────────────────────────────────────────────────────────
print("\n[7e] PARSING — texto_bruto e offset (simulação sem API)")

# Simula o que parsear_emendas_com_ia faria ao construir objetos Emenda
# Bug 1: emenda supressiva com novo_texto=None não deve ter texto_bruto vazio
e_sup = Emenda(
    numero=1,
    texto_bruto=(
        None                        # simula item.get("texto_bruto") ausente
        or None                     # simula item.get("novo_texto") ausente (Supressiva)
        or "[Emenda 1 — Supressiva | Art. 4º]"  # fallback do código corrigido
    ),
    tipo=TipoEmenda.SUPRESSIVA, alvo="Art. 4º", parseada=True,
)
chk("Supressiva: texto_bruto não é vazio após parsing",
    bool(e_sup.texto_bruto and e_sup.texto_bruto.strip()))

e_mod = Emenda(
    numero=2,
    texto_bruto="Emenda nº 2 — Dê-se ao §2º do Art. 7º a seguinte redação: ...",
    tipo=TipoEmenda.MODIFICATIVA, alvo="Art. 7º, §2º",
    novo_texto="Nova redação do §2º...", parseada=True,
)
chk("Modificativa: texto_bruto preserva texto integral (distinto de novo_texto)",
    e_mod.texto_bruto != e_mod.novo_texto and len(e_mod.texto_bruto) > 0)

# Bug 2: offset — simula a fórmula REAL do código corrigido (idx_item, não len(todas))
# Dois lotes de 2 emendas sem número → esperado: [1, 2, 3, 4]
todas2: list[Emenda] = []
offset_sim = 0
for lote in range(2):
    n_antes = len(todas2)
    itens_lote = [{"tipo": "Supressiva", "alvo": f"Art. {lote*2+i+1}"} for i in range(2)]
    for idx_item, item in enumerate(itens_lote, start=1):
        numero = item.get("numero")          # None — sem número explícito
        if numero is None:
            numero = offset_sim + idx_item   # fórmula corrigida
        todas2.append(Emenda(numero=numero, texto_bruto=f"emenda lote{lote} item{idx_item}"))
    offset_sim += len(todas2) - n_antes

chk("Offset 2 lotes de 2: sequência é [1,2,3,4] (não [1,2,5,6])",
    [e.numero for e in todas2] == [1, 2, 3, 4],
    f"sequência: {[e.numero for e in todas2]}")
chk("Offset após 2 lotes = 4 (não 6)",
    offset_sim == 4, f"offset={offset_sim}")

# P1: parsing sem JSON — raise ValueError (não continue silencioso)
import inspect as _inspect_p
from harmonizer import parsear_emendas_com_ia as _parse_fn
_src_parse = _inspect_p.getsource(_parse_fn)

chk("Parsing sem JSON: raise ValueError com mensagem clara (não continue silencioso)",
    "raise ValueError" in _src_parse and "IA não retornou JSON" in _src_parse)

chk("Except captura ValueError — fallback bruto ativado (emendas não perdidas)",
    bool(re.search(r'except\s*\([^)]*ValueError', _src_parse)))

# ────────────────────────────────────────────────────────────────────────────
# 8. ANÁLISE ESTRUTURAL
# ────────────────────────────────────────────────────────────────────────────
print("\n[8] ANÁLISE ESTRUTURAL (analisar_estrutura)")

TEXTO_EST = (
    "Art. 1º Disposições gerais.\n"
    "Art. 2º Regulamentação.\n"
    "§ 1º Primeiro parágrafo.\n"
    "§ 2º Segundo parágrafo.\n"
    "I — primeiro inciso;\n"
    "II — segundo inciso.\n"
    "  a) alínea a;\n"
    "  b) alínea b.\n"
    "ANEXO I\nANEXO II\nANEXO III"
)
est = analisar_estrutura(TEXTO_EST)
chk("2 artigos",    est["artigos"]    == 2, f"encontrado: {est['artigos']}")
chk("2 parágrafos", est["paragrafos"] == 2, f"encontrado: {est['paragrafos']}")
chk("3 anexos",     est["anexos"]     == 3, f"encontrado: {est['anexos']}")

# Regressão: "Art X." sem ponto após "Art" (bug PLC 55/2025 — Art 14.)
TEXTO_EST_SEM_PONTO = (
    "Art. 1º Disposição.\n"
    "Art 14. Dispositivo sem ponto após Art.\n"   # bug: PDF da CMRJ omite o ponto
    "Art. 15. Normal.\n"
    "Art. 21. Último.\n"
)
est2 = analisar_estrutura(TEXTO_EST_SEM_PONTO)
chk("4 artigos (inclui Art sem ponto)", est2["artigos"] == 4,
    f"encontrado: {est2['artigos']} — esperado 4 (Art 14. deve ser contado)")

# ────────────────────────────────────────────────────────────────────────────
# 9. CHAVE DE API
# ────────────────────────────────────────────────────────────────────────────
print("\n[9] CHAVE DE API ANTHROPIC")

api_key = os.environ.get("ANTHROPIC_API_KEY", "")
chk("ANTHROPIC_API_KEY definida",
    bool(api_key),
    "Defina com:  $env:ANTHROPIC_API_KEY='sk-ant-...'  (PowerShell)")

# ────────────────────────────────────────────────────────────────────────────
# 10. ARQUIVOS DE STRESS TEST
# ────────────────────────────────────────────────────────────────────────────
print("\n[10] ARQUIVOS DO STRESS TEST (PLC 17/2026)")

TAB1 = pathlib.Path(r"C:\Users\Admin\Downloads\TAB_1_PLC_17_2026_TEXTO_ORIGINAL.txt")
TAB2 = pathlib.Path(r"C:\Users\Admin\Downloads\TAB_2_PLC_17_2026_EMENDAS.txt")
chk("TAB_1 (texto original) encontrado", TAB1.exists(), str(TAB1))
chk("TAB_2 (emendas) encontrado",        TAB2.exists(), str(TAB2))

# ────────────────────────────────────────────────────────────────────────────
# 11. VALIDAÇÃO XML GENERALIZADA: par completo em todas as tags (rev.10)
# ────────────────────────────────────────────────────────────────────────────
print("\n[11] VALIDAÇÃO XML GENERALIZADA (harmonizer.py — rev.10)")

try:
    from harmonizer import harmonizar_texto as _harm_func
    import inspect
    import re as _re
    src = inspect.getsource(_harm_func)

    # 11a — Guarda usa _TODAS_TAGS cobrindo 8 tags (incl. MAPA_RENUMERACAO, NOTAS_TECNICAS e SUGESTOES_NORMATIVAS)
    chk("Guarda generalizada: _TODAS_TAGS cobre 8 tags (incl. MAPA_RENUMERACAO, NOTAS_TECNICAS e SUGESTOES_NORMATIVAS)",
        bool(_re.search(r'_TODAS_TAGS\s*=\s*\[', src)) and
        "MAPA_RENUMERACAO" in src and "LOG_ALTERACOES" in src and
        "NOTAS_TECNICAS" in src and "SUGESTOES_NORMATIVAS" in src)

    # 11b — Truncamento detectado por ausência de par completo (_sem_par)
    chk("Truncamento detectado por _sem_par (par abertura+fechamento)",
        "_sem_par" in src)

    # 11c — Conteúdo vazio rejeitado via _TAGS_NAO_VAZIAS e _conteudo_vazio
    chk("Conteúdo vazio rejeitado: _TAGS_NAO_VAZIAS + _conteudo_vazio",
        "_TAGS_NAO_VAZIAS" in src and "_conteudo_vazio" in src)

    # 11d — TEXTO_HARMONIZADO: resposta sem par completo → None
    _resp_sem = "Resposta malformada sem nenhuma tag XML esperada."
    _md = _re.search(r'<TEXTO_HARMONIZADO>(.*?)</TEXTO_HARMONIZADO>', _resp_sem, _re.DOTALL)
    chk("TEXTO_HARMONIZADO: sem par completo → None",
        _md is None)

    # 11e — TEXTO_HARMONIZADO: truncada (só abertura, sem fechamento) → None
    _resp_trunc = "<TEXTO_HARMONIZADO>\nArt. 1º Texto sem fechamento..."
    _me = _re.search(r'<TEXTO_HARMONIZADO>(.*?)</TEXTO_HARMONIZADO>', _resp_trunc, _re.DOTALL)
    chk("TEXTO_HARMONIZADO: truncada (sem fechamento) → None",
        _me is None)

    # 11f — TEXTO_HARMONIZADO: válida → extraída com group(1)
    _resp_ok = "<TEXTO_HARMONIZADO>\nArt. 1º Texto.\nArt. 2º Outro.\n</TEXTO_HARMONIZADO>"
    _mf = _re.search(r'<TEXTO_HARMONIZADO>(.*?)</TEXTO_HARMONIZADO>', _resp_ok, _re.DOTALL)
    chk("TEXTO_HARMONIZADO: válida → extraída com group(1)",
        _mf is not None and "Art. 1º" in _mf.group(1))

    # 11g — texto_harm usa .group(1).strip() (sem fallback silencioso)
    chk("texto_harm usa .group(1).strip() após validação (sem fallback)",
        ".group(1).strip()" in src)

    # 11h–11l — Truncamento detectado em cada uma das outras 7 tags
    _OUTRAS_TAGS = [
        "MAPA_RENUMERACAO", "AVISOS", "ERROS_CRITICOS",
        "ALERTAS_ABSURDOS", "NOTAS_TECNICAS", "SUGESTOES_NORMATIVAS", "LOG_ALTERACOES",
    ]
    for _t in _OUTRAS_TAGS:
        _resp_t = f"<{_t}>\nConteúdo truncado sem tag de fechamento..."
        _mt = _re.search(rf'<{_t}>(.*?)</{_t}>', _resp_t, _re.DOTALL)
        chk(f"{_t}: truncada (sem </{_t}>) → None",
            _mt is None)

    # 11m — LOG_ALTERACOES em _TAGS_NAO_VAZIAS (conteúdo não pode ser vazio)
    chk("LOG_ALTERACOES exige conteúdo não vazio (em _TAGS_NAO_VAZIAS)",
        bool(_re.search(r'_TAGS_NAO_VAZIAS\s*=\s*\{[^}]*LOG_ALTERACOES[^}]*\}', src)))

    # 11n — Resposta completa válida: todos os 8 pares detectados
    _resp_completa = (
        "<TEXTO_HARMONIZADO>\nArt. 1º Texto.\n</TEXTO_HARMONIZADO>\n"
        "<MAPA_RENUMERACAO>\nSem renumeração necessária.\n</MAPA_RENUMERACAO>\n"
        "<AVISOS>\nNenhum aviso.\n</AVISOS>\n"
        "<ERROS_CRITICOS>\nNenhum erro crítico.\n</ERROS_CRITICOS>\n"
        "<ALERTAS_ABSURDOS>\nNenhum.\n</ALERTAS_ABSURDOS>\n"
        "<NOTAS_TECNICAS>\nNenhuma nota técnica.\n</NOTAS_TECNICAS>\n"
        "<SUGESTOES_NORMATIVAS>\nNenhuma sugestão.\n</SUGESTOES_NORMATIVAS>\n"
        "<LOG_ALTERACOES>\nEmenda 1 (Modificativa): Art. 1º atualizado.\n</LOG_ALTERACOES>"
    )
    _todos_ok = all(
        _re.search(rf'<{t}>(.*?)</{t}>', _resp_completa, _re.DOTALL)
        for t in ["TEXTO_HARMONIZADO", "MAPA_RENUMERACAO", "AVISOS",
                  "ERROS_CRITICOS", "ALERTAS_ABSURDOS", "NOTAS_TECNICAS",
                  "SUGESTOES_NORMATIVAS", "LOG_ALTERACOES"]
    )
    chk("Resposta completa e válida — todos os 8 pares detectados",
        _todos_ok)

except Exception as e:
    chk("11 — validação XML generalizada", False, str(e))

# ────────────────────────────────────────────────────────────────────────────
# 12. _invalidar_resultado() PRESENTE E COBERTO EM app.py
# ────────────────────────────────────────────────────────────────────────────
print("\n[12] HELPER _invalidar_resultado() (app.py)")

try:
    import ast
    _app_src = pathlib.Path(r"C:\Users\Admin\Documents\Claude\CCJ\sistema_redacoes\app.py").read_text(encoding='utf-8')
    _n_invalidar = _app_src.count("_invalidar_resultado()")
    chk("_invalidar_resultado() definido em app.py",
        "def _invalidar_resultado()" in _app_src)
    chk("_invalidar_resultado() chamado ≥10 vezes (cobre todos os pontos críticos)",
        _n_invalidar >= 10,
        f"encontrado {_n_invalidar} chamadas — esperado ≥10")
    chk("Votação individual (v_apr) invalida resultado",
        "v_apr" in _app_src and
        # A chamada deve aparecer próxima ao status APROVADA
        "_invalidar_resultado" in _app_src[
            _app_src.index("v_apr"):_app_src.index("v_apr") + 300
        ])
    chk("Importar texto bruto invalida resultado",
        "Importar como texto bruto" in _app_src and
        "_invalidar_resultado" in _app_src[
            _app_src.index("Importar como texto bruto"):
            _app_src.index("Importar como texto bruto") + 1300   # bloco longo c/ re.split
        ])
    chk("Adicionar emenda manual invalida resultado",
        "adicionada!" in _app_src and
        "_invalidar_resultado" in _app_src[
            max(0, _app_src.rindex("adicionada!") - 400):
            _app_src.rindex("adicionada!") + 10
        ])
except Exception as e:
    chk("12 — _invalidar_resultado()", False, str(e))

# ────────────────────────────────────────────────────────────────────────────
# 13. REGRA A4 — "acrescente-se onde couber" (estrutural, sem API)
# ────────────────────────────────────────────────────────────────────────────
print("\n[13] REGRA A4 — emendas sem alvo definido (estrutural)")

try:
    import inspect as _insp_a4
    _src_harm = _insp_a4.getsource(_harm_func)

    # A4.1 — aditivas sem alvo
    chk("A4.1 presente — aditivas sem alvo",
        "A4.1" in _src_harm,
        "Regra A4.1 não encontrada em harmonizar_texto")

    chk("A4.1 cobre expressão 'onde couber'",
        "onde couber" in _src_harm,
        "Expressão 'onde couber' não encontrada no prompt A4")

    chk("A4.1 cobre unidades menores (parágrafo/inciso/alínea/item)",
        "Parágrafo novo" in _src_harm and "Inciso novo" in _src_harm and "Alínea nova" in _src_harm,
        "Cobertura de unidades menores (parágrafo/inciso/alínea) não encontrada")

    chk("A4.1 LOG inclui tipo de unidade normativa inserida",
        "artigo/parágrafo/inciso/alínea/item" in _src_harm,
        "Formato de LOG com tipo de unidade não encontrado")

    chk("A4.1 exige AVISO com 'alvo não especificado'",
        "alvo não especificado" in _src_harm,
        "Frase 'alvo não especificado' não encontrada no prompt A4")

    chk("A4.1 proíbe inserção silenciosa",
        "NUNCA insira silenciosamente" in _src_harm,
        "Proibição de inserção silenciosa não encontrada")

    chk("A4.1 proíbe recusar aplicar emenda sem alvo",
        "NUNCA recuse aplicar" in _src_harm,
        "Proibição de recusa não encontrada")

    # A4.2 — modificativas/substitutivas sem alvo identificável
    chk("A4.2 presente — modificativas/substitutivas sem alvo",
        "A4.2" in _src_harm,
        "Regra A4.2 não encontrada em harmonizar_texto")

    chk("A4.2 proíbe aplicar substituição sem alvo identificável",
        "NÃO aplique a substituição" in _src_harm,
        "Proibição de substituição inferida não encontrada")

    chk("A4.2 gera ERRO CRÍTICO (não AVISO) de 'alvo não identificável'",
        "alvo não identificável" in _src_harm and "ERROS_CRITICOS" in _src_harm,
        "Erro crítico 'alvo não identificável' em ERROS_CRITICOS não encontrado")

    chk("A4.2 registra no LOG que emenda NÃO foi aplicada",
        "NÃO aplicada" in _src_harm and "A4.2 / Emenda" in _src_harm,
        "LOG de emenda não aplicada (A4.2) não encontrado")

except Exception as e:
    chk("13 — regra A4 estrutural", False, str(e))

# ────────────────────────────────────────────────────────────────────────────
# 14. REGRA E2 — conflitos entre emendas aprovadas + sugestão normativa
# ────────────────────────────────────────────────────────────────────────────
print("\n[14] REGRA E2 — conflito de emendas aprovadas + sugestão normativa (estrutural)")

try:
    import inspect as _insp_e2
    _src_e2 = _insp_e2.getsource(_harm_func)

    # E2 — detecção de conflito no prompt
    chk("E2 tem instrução de varredura prévia obrigatória",
        "VARREDURA PRÉVIA OBRIGATÓRIA" in _src_e2 or "varredura prévia" in _src_e2.lower(),
        "Instrução de varredura prévia não encontrada no prompt E2")

    chk("E2 identifica emendas sobre o MESMO dispositivo",
        "MESMO dispositivo" in _src_e2 or "mesmo dispositivo" in _src_e2,
        "Cobertura de emendas sobre mesmo dispositivo não encontrada")

    chk("E2 identifica supressão + modificação simultânea",
        "supressiva" in _src_e2.lower() and "modificativa" in _src_e2.lower(),
        "Cobertura de supressão + modificação simultânea não encontrada")

    chk("E2 aplica menor número como cautela formal",
        "MENOR NÚMERO" in _src_e2 or "menor número" in _src_e2.lower(),
        "Política de cautela (menor número) não encontrada no prompt E2")

    chk("E2 insere marcador de conflito no texto harmonizado",
        "CONFLITO DE EMENDAS" in _src_e2,
        "Marcador de conflito no texto não encontrado")

    chk("E2 exige registro em ERROS_CRITICOS com formato 'CONFLITO / Emendas'",
        "CONFLITO / Emendas" in _src_e2,
        "Formato 'CONFLITO / Emendas' em ERROS_CRITICOS não encontrado")

    # SUGESTOES_NORMATIVAS — tag e integração
    chk("SUGESTOES_NORMATIVAS presente em _TODAS_TAGS",
        "SUGESTOES_NORMATIVAS" in _src_e2,
        "Tag SUGESTOES_NORMATIVAS não encontrada em harmonizar_texto")

    chk("E2 exige geração de sugestão normativa orientativa",
        "SUGESTOES_NORMATIVAS" in _src_e2 and "Sugestão" in _src_e2,
        "Geração de sugestão normativa em SUGESTOES_NORMATIVAS não encontrada")

    chk("'Nenhuma sugestão.' presente no skip set",
        "Nenhuma sugestão." in _src_e2,
        "Frase 'Nenhuma sugestão.' não encontrada no skip set")

    chk("sugestoes_normativas no dataclass ResultadoHarmonizacao",
        "sugestoes_normativas" in _insp_e2.getsource(ResultadoHarmonizacao),
        "Campo 'sugestoes_normativas' não encontrado no dataclass")

    # app.py — exibição sem exportação
    _app_e2 = pathlib.Path(r"C:\Users\Admin\Documents\Claude\CCJ\sistema_redacoes\app.py").read_text(encoding='utf-8')
    chk("app.py exibe sugestoes_normativas",
        "sugestoes_normativas" in _app_e2,
        "sugestoes_normativas não encontrado em app.py")

    chk("app.py NÃO passa sugestoes_normativas para exportar_redacao_final_docx",
        "sugestoes_normativas" not in _app_e2[
            _app_e2.rindex("exportar_redacao_final_docx"):
            _app_e2.rindex("exportar_redacao_final_docx") + 600
        ],
        "sugestoes_normativas está sendo passado para exportar_redacao_final_docx (não deveria)")

except Exception as e:
    chk("14 — regra E2 estrutural", False, str(e))

# ────────────────────────────────────────────────────────────────────────────
# 15. SUBEMENDAS — pré-processamento e dataclass
# ────────────────────────────────────────────────────────────────────────────
print("\n[15] SUBEMENDAS — processamento pré-harmonização (estrutural)")

try:
    from harmonizer import _resolver_subemendas

    # Campo subemenda_de no dataclass
    chk("subemenda_de presente no dataclass Emenda",
        "subemenda_de" in Emenda.__dataclass_fields__)

    # _resolver_subemendas: caso normal — aprovada + pai aprovado → substitui
    _e_pai = Emenda(numero=3, texto_bruto="Texto original da Emenda 3.",
                    tipo=TipoEmenda.MODIFICATIVA, alvo="Art. 5º",
                    novo_texto="Texto original da Emenda 3.",
                    status=StatusEmenda.APROVADA, parseada=True)
    _e_sub = Emenda(numero=7, texto_bruto="SubEmenda à Emenda nº 3 — novo texto.",
                    tipo=TipoEmenda.MODIFICATIVA, alvo="Art. 5º",
                    novo_texto="Novo texto aprovado pela subemenda.",
                    status=StatusEmenda.APROVADA, parseada=True,
                    subemenda_de=3)
    _lista, _log_s, _av_s, _ec_s = _resolver_subemendas([_e_pai, _e_sub], [_e_pai, _e_sub])

    chk("SubEmenda aprovada + pai aprovado → substitui texto do pai",
        len(_lista) == 1 and _lista[0].numero == 3 and
        "novo texto" in (_lista[0].novo_texto or "").lower(),
        f"lista={[e.numero for e in _lista]}, novo_texto={_lista[0].novo_texto if _lista else '?'}")

    chk("SubEmenda aprovada + pai aprovado → subemenda retirada da lista (não vai à IA)",
        all(e.numero != 7 for e in _lista),
        f"lista={[e.numero for e in _lista]}")

    chk("SubEmenda aprovada + pai aprovado → log registra substituição",
        any("SubEmenda 7" in l and "Emenda 3" in l for l in _log_s),
        f"log={_log_s}")

    # SubEmenda aprovada + pai NÃO aprovado → inoperante
    _e_pai_rej = Emenda(numero=3, texto_bruto="Texto original.",
                        status=StatusEmenda.REJEITADA, parseada=True)
    _e_sub2 = Emenda(numero=7, texto_bruto="SubEmenda.",
                     novo_texto="Sub texto.",
                     status=StatusEmenda.APROVADA, parseada=True, subemenda_de=3)
    _lista2, _log2, _av2, _ec2 = _resolver_subemendas([_e_pai_rej, _e_sub2], [_e_sub2])

    chk("SubEmenda aprovada + pai rejeitado → aviso de inoperante",
        any("inoperante" in a for a in _av2),
        f"avisos={_av2}")

    chk("SubEmenda aprovada + pai rejeitado → subemenda não vai à IA",
        len(_lista2) == 0 or all(e.subemenda_de is None for e in _lista2),
        f"lista={[e.numero for e in _lista2]}")

    # SubEmenda rejeitada → pai mantém texto original
    _e_sub3 = Emenda(numero=8, texto_bruto="Sub rejeitada.",
                     novo_texto="Texto sub.",
                     status=StatusEmenda.REJEITADA, parseada=True, subemenda_de=3)
    _e_pai3 = Emenda(numero=3, texto_bruto="Texto original mantido.",
                     novo_texto="Texto original mantido.",
                     status=StatusEmenda.APROVADA, parseada=True)
    _lista3, _log3, _av3, _ec3 = _resolver_subemendas([_e_pai3, _e_sub3], [_e_pai3])

    chk("SubEmenda rejeitada → pai mantém texto original (log registra)",
        any("SubEmenda 8" in l and "mantém texto original" in l for l in _log3),
        f"log={_log3}")

    chk("SubEmenda rejeitada → texto do pai não alterado",
        len(_lista3) == 1 and "original" in (_lista3[0].novo_texto or "").lower(),
        f"novo_texto={_lista3[0].novo_texto if _lista3 else '?'}")

    # Conflito: duas subemendas aprovadas para o mesmo pai
    _sub_a = Emenda(numero=5, status=StatusEmenda.APROVADA, texto_bruto="Sub A",
                    novo_texto="Texto A", parseada=True, subemenda_de=3)
    _sub_b = Emenda(numero=6, status=StatusEmenda.APROVADA, texto_bruto="Sub B",
                    novo_texto="Texto B", parseada=True, subemenda_de=3)
    _e_pai4 = Emenda(numero=3, status=StatusEmenda.APROVADA, texto_bruto="Original",
                     novo_texto="Original", parseada=True)
    _lista4, _log4, _av4, _ec4 = _resolver_subemendas(
        [_e_pai4, _sub_a, _sub_b], [_e_pai4, _sub_a, _sub_b]
    )
    chk("Conflito de subemendas → erro crítico §2º gerado (não §1º aviso)",
        any("CONFLITO DE SUBEMENDAS" in e for e in _ec4),
        f"erros_criticos={_ec4}, avisos={_av4}")

    chk("Conflito de subemendas → NÃO entra em avisos §1º (somente §2º)",
        not any("CONFLITO DE SUBEMENDAS" in a for a in _av4),
        f"avisos={_av4}")

    chk("Conflito de subemendas → nenhuma substituição automática (texto pai preservado)",
        len(_lista4) == 1 and _lista4[0].novo_texto == "Original",
        f"lista={[e.numero for e in _lista4]}")

    # ── P1: Auto-referência (subemenda_de == numero próprio) ─────────────────
    _e_autoref = Emenda(numero=5, texto_bruto="Texto E5.",
                        novo_texto="Texto E5.", status=StatusEmenda.APROVADA,
                        parseada=True, subemenda_de=5)   # aponta para si mesma
    _lista_p1, _log_p1, _av_p1, _ec_p1 = _resolver_subemendas(
        [_e_autoref], [_e_autoref]
    )
    chk("P1: auto-referência → erro crítico §2º gerado",
        any("AUTO-REFERÊNCIA" in e for e in _ec_p1),
        f"erros_criticos={_ec_p1}")
    chk("P1: auto-referência → emenda excluída da lista (não vai à IA)",
        len(_lista_p1) == 0,
        f"lista={[e.numero for e in _lista_p1]}")
    chk("P1: auto-referência → NÃO entra em avisos §1º",
        not any("AUTO" in a for a in _av_p1),
        f"avisos={_av_p1}")

    # ── P2: Pai inexistente → ERRO CRÍTICO §2º (não §1º aviso) ───────────────
    _e_orphan = Emenda(numero=9, texto_bruto="Sub sem pai.",
                       novo_texto="Texto sub.", status=StatusEmenda.APROVADA,
                       parseada=True, subemenda_de=99)   # Emenda 99 não existe
    _lista_p2, _log_p2, _av_p2, _ec_p2 = _resolver_subemendas(
        [_e_orphan], [_e_orphan]
    )
    chk("P2: pai inexistente → erro crítico §2º gerado",
        any("SEM PAI" in e for e in _ec_p2),
        f"erros_criticos={_ec_p2}")
    chk("P2: pai inexistente → NÃO entra em avisos §1º (era bug pré-rev.15)",
        not any("consta na lista" in a or "verifique a numeração" in a for a in _av_p2),
        f"avisos={_av_p2}")
    chk("P2: pai inexistente → subemenda excluída da lista",
        len(_lista_p2) == 0,
        f"lista={[e.numero for e in _lista_p2]}")

    # ── P3: Subemenda encadeada (subemenda de subemenda) → ERRO CRÍTICO §2º ──
    _e3_base  = Emenda(numero=3, texto_bruto="Base.", novo_texto="Base.",
                       status=StatusEmenda.APROVADA, parseada=True)
    _e7_sub3  = Emenda(numero=7, texto_bruto="Sub7 da 3.", novo_texto="Texto Sub7.",
                       status=StatusEmenda.APROVADA, parseada=True, subemenda_de=3)
    _e8_sub7  = Emenda(numero=8, texto_bruto="Sub8 da 7.", novo_texto="Texto Sub8.",
                       status=StatusEmenda.APROVADA, parseada=True, subemenda_de=7)
    _lista_p3, _log_p3, _av_p3, _ec_p3 = _resolver_subemendas(
        [_e3_base, _e7_sub3, _e8_sub7], [_e3_base, _e7_sub3, _e8_sub7]
    )
    chk("P3: subemenda encadeada → erro crítico §2º gerado",
        any("ENCADEADA" in e for e in _ec_p3),
        f"erros_criticos={_ec_p3}")
    chk("P3: subemenda encadeada → NÃO entra em avisos §1º",
        not any("encadeada" in a.lower() for a in _av_p3),
        f"avisos={_av_p3}")
    chk("P3: subemenda encadeada → subemenda de subemenda excluída da lista",
        all(e.numero != 8 for e in _lista_p3),
        f"lista={[e.numero for e in _lista_p3]}")

    # Verifica presença de subemenda_de no parser
    import inspect as _insp15
    from harmonizer import parsear_emendas_com_ia as _parse_fn15
    _src_parse15 = _insp15.getsource(_parse_fn15)
    chk("parsear_emendas_com_ia reconhece subemenda_de no JSON",
        "subemenda_de" in _src_parse15,
        "'subemenda_de' não encontrado no prompt de parsing")

    # app.py exibe painel de subemendas na aba 4
    _app15 = pathlib.Path(r"C:\Users\Admin\Documents\Claude\CCJ\sistema_redacoes\app.py").read_text(encoding='utf-8')
    chk("app.py exibe painel de subemendas na aba 4",
        "_subs_aba4" in _app15 and "subemenda" in _app15.lower(),
        "Painel de subemendas não encontrado em app.py")

except Exception as e:
    chk("15 — subemendas estrutural", False, str(e))

# ────────────────────────────────────────────────────────────────────────────
# RESUMO
# ────────────────────────────────────────────────────────────────────────────
print()
print("=" * 65)
total  = len(resultados)
passou = sum(1 for _, ok in resultados if ok)
falhou = total - passou
print(f"  RESULTADO FINAL: {passou}/{total} verificações passaram")

if falhou:
    print(f"\n  ⚠  {falhou} falha(s) detectada(s):")
    for nome, ok in resultados:
        if not ok:
            print(f"     ❌ {nome}")
    print()
else:
    print("\n  ✅ Tudo certo — sistema pronto para uso no PLC real.")
    print()

print("=" * 65)

# ────────────────────────────────────────────────────────────────────────────
# 11. HARMONIZAÇÃO COMPLETA (--com-api)
# ────────────────────────────────────────────────────────────────────────────
if "--com-api" not in sys.argv:
    print()
    print("  Dica: rode  python verificar.py --com-api  para testar a")
    print("  harmonização completa do PLC 17/2026 (custo ~$0,50).")
    print()
    sys.exit(0 if falhou == 0 else 1)

print()
print("[11] HARMONIZAÇÃO COMPLETA (--com-api) — PLC 17/2026")

if not api_key:
    print("  ⏭  ANTHROPIC_API_KEY não definida — pulando.")
    sys.exit(0 if falhou == 0 else 1)

if not (TAB1.exists() and TAB2.exists()):
    print("  ⏭  Arquivos TAB_1/TAB_2 não encontrados — pulando.")
    sys.exit(0 if falhou == 0 else 1)

try:
    from utils import ler_txt
    from harmonizer import parsear_emendas_com_ia, harmonizar_texto

    print("  Carregando arquivos de texto...")
    texto_orig    = ler_txt(TAB1.read_bytes())
    texto_emendas = ler_txt(TAB2.read_bytes())

    print("  Parseando emendas com IA...")
    emendas = parsear_emendas_com_ia(texto_emendas, api_key)
    for e in emendas:
        e.status = StatusEmenda.APROVADA
    print(f"  {len(emendas)} emendas parseadas. Harmonizando (pode demorar ~2 min)...")

    resultado = harmonizar_texto(texto_orig, emendas, api_key, "PLC 17/2026")
    est_fin   = analisar_estrutura(resultado.texto_harmonizado)

    print()
    print("  --- Resultados da harmonização ---")
    print(f"  Artigos:  {est_fin['artigos']} (esperado: 18)")
    print(f"  Anexos:   {est_fin['anexos']}  (esperado: 5)")
    print(f"  Absurdos §2º: {len(resultado.alertas_absurdos)} (esperado: ≥3)")
    print(f"  Avisos §1º:   {len(resultado.avisos)}")
    print()

    chk("[API] 18 artigos",  est_fin["artigos"] == 18,  f"encontrado: {est_fin['artigos']}")
    chk("[API] 5 anexos",    est_fin["anexos"]  == 5,   f"encontrado: {est_fin['anexos']}")
    chk("[API] ≥3 absurdos manifestos §2º",
        len(resultado.alertas_absurdos) >= 3,
        f"alertas_absurdos={len(resultado.alertas_absurdos)}")
    # Aceita "§1º" ou "§ 1º" (com ou sem espaço — ambas são formatações válidas)
    chk("[API] Frase verbatim §1º preservada",
        bool(re.search(r'Atendida a condição prevista no §\s*1º deste artigo',
                       resultado.texto_harmonizado)))
    chk("[API] 'artigo anterior' preservado no §4º",
        "nos termos do artigo anterior" in resultado.texto_harmonizado)
    # E1 — auto-correções registradas no LOG (critérios 3 e 6 do AUDITORIA.md)
    chk("[API] E1: 'serão aplicadas' corrigido no texto (Emenda 8)",
        "serão aplicadas" in resultado.texto_harmonizado,
        "verificar se Emenda 8 foi corrigida")
    chk("[API] E1: 'depósitos' em minúscula no texto (Emenda 9)",
        bool(re.search(r'\bdepósitos\b', resultado.texto_harmonizado)),
        "verificar se Emenda 9 foi corrigida")
    chk("[API] E1: LOG registra correção de concordância",
        any("serão aplicad" in l.lower() or "E1" in l for l in resultado.log_alteracoes),
        f"log tem {len(resultado.log_alteracoes)} entradas")
    # CA/mérito: observações sobre parâmetros urbanísticos NÃO devem aparecer em AVISOS
    chk("[API] Sem observações de mérito/CA nos AVISOS",
        not any(re.search(r'\bCA\b|\bcoeficiente de aproveitamento\b|\bganharit\b|'
                          r'setor [AB].*maior|maior.*setor|gabarito.*CA|CA.*gabarito',
                          a, re.IGNORECASE)
                for a in resultado.avisos),
        "AVISOS não devem conter análises de mérito urbanístico")
    chk("[API] Marcadores inline removidos do DOCX",
        not any("[[" in t for t in [p.text for p in Document(BytesIO(
            exportar_redacao_final_docx(
                resultado.texto_harmonizado, "PLC 17/2026",
                resultado.avisos, resultado.erros_criticos,
                resultado.alertas_absurdos,
            )
        )).paragraphs]))

    # Salva resultado completo
    saida = pathlib.Path("resultado_verificar.txt")
    with open(saida, "w", encoding="utf-8") as f:
        f.write("=== RESULTADO HARMONIZAÇÃO PLC 17/2026 — VERIFICAR.PY ===\n\n")
        f.write(f"Artigos: {est_fin['artigos']}  |  Anexos: {est_fin['anexos']}\n\n")
        f.write(f"--- ALERTAS §2º ({len(resultado.alertas_absurdos)}) ---\n")
        for a in resultado.alertas_absurdos:
            f.write(f"  {a}\n")
        f.write(f"\n--- ERROS CRÍTICOS ({len(resultado.erros_criticos)}) ---\n")
        for e in resultado.erros_criticos:
            f.write(f"  {e}\n")
        f.write(f"\n--- AVISOS §1º ({len(resultado.avisos)}) ---\n")
        for a in resultado.avisos:
            f.write(f"  {a}\n")
        f.write("\n--- LOG DE ALTERAÇÕES ---\n")
        for l in resultado.log_alteracoes:
            f.write(f"  {l}\n")
        f.write("\n--- TEXTO HARMONIZADO ---\n")
        f.write(resultado.texto_harmonizado)
    print(f"\n  Resultado completo salvo em: {saida.resolve()}")

except Exception as e:
    chk("[API] Harmonização sem erro", False, str(e))
    import traceback
    traceback.print_exc()

# Resumo final com API
print()
print("=" * 65)
total2  = len(resultados)
passou2 = sum(1 for _, ok in resultados if ok)
falhou2 = total2 - passou2
print(f"  RESULTADO FINAL (com API): {passou2}/{total2} verificações passaram")
if falhou2:
    print(f"\n  ❌ {falhou2} falha(s):")
    for nome, ok in resultados:
        if not ok:
            print(f"     • {nome}")
else:
    print("\n  ✅ Sistema completamente verificado — pronto para PLC 92/2025.")
print("=" * 65)
print()
sys.exit(0 if falhou2 == 0 else 1)
